from openai import OpenAI
from docx import Document

import os
from openai import OpenAI

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

print("\n=== AI FUTURE BUSINESS + CONTENT GENERATOR ===\n")

business_type = input("What business or content do you do? ")
audience = input("Who is your target audience? ")
platform = input("Which platform(s)? ")

print("\nGenerating FULL Business + Content Kit...\n")

response = client.chat.completions.create(
    model="gpt-4.1-mini",
    messages=[
        {
            "role": "user",
            "content": f"""
Create a FULL BUSINESS + CONTENT KIT.

Business Type: {business_type}
Audience: {audience}
Platform: {platform}

Brand Style:
History + AI + Money + Future Tech + Education + Motivation

Generate:

SECTION 1 — 30 Viral Content Ideas

SECTION 2 — 10 Digital Product Ideas

SECTION 3 — 5 Short Video Scripts

SECTION 4 — Marketing Strategy Plan

SECTION 5 — Monetization Strategy Plan
"""
        }
    ],
    max_tokens=1500
)

result_text = response.choices[0].message.content

print("\n=== FULL AI RESULT ===\n")
print(result_text)

# Save TXT
with open("AI_Full_Business_Kit.txt", "w", encoding="utf-8") as f:
    f.write(result_text)

# Save Word
doc = Document()
doc.add_heading("AI Full Business + Content Kit", level=1)
doc.add_paragraph(result_text)
doc.save("AI_Full_Business_Kit.docx")

print("\n✅ Full Business Kit Saved (TXT + Word)!")
